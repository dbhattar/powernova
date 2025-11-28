"""
Document processor service
Extracts text content from various document formats
"""
import logging
from typing import Tuple, Optional
from bs4 import BeautifulSoup
import PyPDF2
from io import BytesIO
import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process and extract text from various document formats
    """
    
    @staticmethod
    def extract_text_from_html(html_content: bytes, url: str) -> Tuple[str, str, dict]:
        """
        Extract text and metadata from HTML
        
        Args:
            html_content: HTML content as bytes
            url: URL of the page
            
        Returns:
            Tuple of (title, extracted_text, metadata)
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extract title
            title = ""
            if soup.title:
                title = soup.title.string.strip() if soup.title.string else ""
            elif soup.find('h1'):
                title = soup.find('h1').get_text().strip()
            else:
                title = url
            
            # Extract metadata
            metadata = {}
            
            # Meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc and meta_desc.get('content'):
                metadata['description'] = meta_desc['content']
            
            # Meta keywords
            meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
            if meta_keywords and meta_keywords.get('content'):
                metadata['keywords'] = meta_keywords['content']
            
            # Author
            meta_author = soup.find('meta', attrs={'name': 'author'})
            if meta_author and meta_author.get('content'):
                metadata['author'] = meta_author['content']
            
            # Extract main text content
            text = soup.get_text(separator=' ', strip=True)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            logger.info(f"Extracted {len(text)} chars from HTML: {title}")
            return title, text, metadata
            
        except Exception as e:
            logger.error(f"Failed to extract text from HTML: {e}")
            return url, "", {}
    
    @staticmethod
    def extract_text_from_pdf(pdf_content: bytes, url: str) -> Tuple[str, str, dict]:
        """
        Extract text and metadata from PDF
        
        Args:
            pdf_content: PDF content as bytes
            url: URL of the PDF
            
        Returns:
            Tuple of (title, extracted_text, metadata)
        """
        title = url
        text = ""
        metadata = {}
        
        try:
            pdf_file = BytesIO(pdf_content)
            
            # Try with strict=False to handle malformed PDFs
            try:
                pdf_reader = PyPDF2.PdfReader(pdf_file, strict=False)
            except Exception as e:
                logger.warning(f"Failed to read PDF with lenient mode: {e}")
                # Try one more time with strict mode
                pdf_file.seek(0)
                pdf_reader = PyPDF2.PdfReader(pdf_file, strict=True)
            
            # Extract metadata (may fail for corrupted PDFs)
            try:
                pdf_info = pdf_reader.metadata
                if pdf_info:
                    if pdf_info.get('/Title'):
                        title = str(pdf_info['/Title'])
                    if pdf_info.get('/Author'):
                        metadata['author'] = str(pdf_info['/Author'])
                    if pdf_info.get('/Subject'):
                        metadata['subject'] = str(pdf_info['/Subject'])
                    if pdf_info.get('/Keywords'):
                        metadata['keywords'] = str(pdf_info['/Keywords'])
            except Exception as e:
                logger.debug(f"Could not extract PDF metadata: {e}")
                metadata['extraction_warning'] = 'Metadata extraction failed'
            
            # Extract text from all pages
            text_parts = []
            failed_pages = 0
            total_pages = 0
            
            try:
                total_pages = len(pdf_reader.pages)
            except Exception as e:
                logger.warning(f"Cannot determine page count: {e}")
                # Try to extract what we can
                total_pages = 0
            
            for page_num in range(total_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    failed_pages += 1
                    logger.debug(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # If we couldn't extract any pages, try alternative method
            if not text_parts and total_pages > 0:
                logger.warning(f"Standard extraction failed for all pages, trying alternative method")
                # Note: Could add alternative extraction methods here (e.g., pdfminer, pdfplumber)
                # For now, we'll just return empty text with warning
                metadata['extraction_warning'] = 'Text extraction failed for all pages'
            
            text = ' '.join(text_parts)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            # Add extraction statistics to metadata
            metadata['page_count'] = total_pages
            if failed_pages > 0:
                metadata['failed_pages'] = failed_pages
                metadata['successful_pages'] = total_pages - failed_pages
                metadata['extraction_warning'] = f'{failed_pages}/{total_pages} pages failed to extract'
            
            if text:
                logger.info(f"Extracted {len(text)} chars from PDF ({total_pages} pages, {failed_pages} failed): {title}")
            else:
                logger.warning(f"No text extracted from PDF ({total_pages} pages): {title}")
                if total_pages > 0:
                    metadata['extraction_warning'] = 'PDF may be image-based (scanned) or corrupted'
            
            return title, text, metadata
            
        except PyPDF2.errors.PdfReadError as e:
            # Specific PDF read errors (EOF marker, encryption, etc.)
            error_msg = str(e)
            logger.error(f"PDF read error for {url}: {error_msg}")
            
            # Categorize the error
            if 'EOF marker not found' in error_msg:
                metadata['error'] = 'Incomplete PDF - EOF marker missing'
                metadata['error_type'] = 'incomplete_pdf'
            elif 'encrypted' in error_msg.lower():
                metadata['error'] = 'Encrypted PDF - password required'
                metadata['error_type'] = 'encrypted_pdf'
            elif 'xref' in error_msg.lower():
                metadata['error'] = 'Corrupted PDF - cross-reference table damaged'
                metadata['error_type'] = 'corrupted_xref'
            else:
                metadata['error'] = f'PDF read error: {error_msg}'
                metadata['error_type'] = 'pdf_read_error'
            
            return title, text, metadata
            
        except Exception as e:
            # Other unexpected errors (including PyCryptodome missing)
            error_msg = str(e)
            logger.error(f"Failed to extract text from PDF {url}: {error_msg}")
            
            # Check for specific error types
            if 'PyCryptodome is required' in error_msg:
                metadata['error'] = 'AES-encrypted PDF requires PyCryptodome library'
                metadata['error_type'] = 'encrypted_aes_pdf'
            else:
                metadata['error'] = f'Extraction failed: {error_msg}'
                metadata['error_type'] = 'unknown_error'
            
            return title, text, metadata
    
    @staticmethod
    def extract_text_from_docx(docx_content: bytes, url: str) -> Tuple[str, str, dict]:
        """
        Extract text from DOCX file
        
        Args:
            docx_content: DOCX content as bytes
            url: URL of the document
            
        Returns:
            Tuple of (title, extracted_text, metadata)
        """
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
            import zipfile
            
            docx_file = BytesIO(docx_content)
            
            # First, check if it's a valid ZIP file (DOCX files are ZIP archives)
            try:
                if not zipfile.is_zipfile(docx_file):
                    logger.warning(f"File is not a ZIP archive, cannot be a valid DOCX: {url}")
                    return url, "", {}
                docx_file.seek(0)  # Reset position after check
            except Exception as e:
                logger.warning(f"Cannot verify ZIP format: {url} - {e}")
                return url, "", {}
            
            # Validate it's actually a DOCX file by checking the magic bytes
            # DOCX files are ZIP archives starting with PK (0x50 0x4B)
            if len(docx_content) < 4 or docx_content[:2] != b'PK':
                logger.warning(f"File does not appear to be a valid DOCX (ZIP) file: {url}")
                return url, "", {}
            
            try:
                doc = Document(docx_file)
            except PackageNotFoundError as e:
                # This happens when the file is a valid ZIP but not a Word document
                # (e.g., theme files, other Office XML components)
                error_msg = str(e)
                if 'themeManager' in error_msg or 'theme' in error_msg.lower():
                    logger.warning(f"File is an Office theme file, not a Word document: {url}")
                else:
                    logger.warning(f"File is not a valid Word document: {url} - {error_msg}")
                return url, "", {}
            except Exception as e:
                # Catch other validation errors from python-docx
                logger.warning(f"Cannot parse as DOCX (may be corrupted or different Office format): {url} - {e}")
                return url, "", {}
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            text = ' '.join(text_parts)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            # Extract title (first heading or first paragraph)
            title = url
            if doc.paragraphs and doc.paragraphs[0].text.strip():
                title = doc.paragraphs[0].text.strip()[:200]
            
            # Extract metadata
            metadata = {}
            try:
                core_props = doc.core_properties
                if core_props.title:
                    title = core_props.title
                    metadata['title'] = core_props.title
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.subject:
                    metadata['subject'] = core_props.subject
                if core_props.keywords:
                    metadata['keywords'] = core_props.keywords
            except Exception as e:
                logger.debug(f"Could not extract DOCX metadata: {e}")
            
            logger.info(f"Extracted {len(text)} chars from DOCX: {title}")
            return title, text, metadata
            
        except ImportError:
            logger.error("python-docx not installed")
            return url, "", {}
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            return url, "", {}
    
    @staticmethod
    def extract_text_from_text(text_content: bytes, url: str) -> Tuple[str, str, dict]:
        """
        Extract text from plain text file
        
        Args:
            text_content: Text content as bytes
            url: URL of the document
            
        Returns:
            Tuple of (title, extracted_text, metadata)
        """
        try:
            # Try to decode as UTF-8, fallback to latin-1
            try:
                text = text_content.decode('utf-8')
            except UnicodeDecodeError:
                text = text_content.decode('latin-1', errors='ignore')
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            # Use first line or URL as title
            title = url
            if text:
                first_line = text.split('\n')[0].strip()
                if first_line and len(first_line) < 200:
                    title = first_line
            
            logger.info(f"Extracted {len(text)} chars from text file: {title}")
            return title, text, {}
            
        except Exception as e:
            logger.error(f"Failed to extract text from text file: {e}")
            return url, "", {}
    
    @staticmethod
    def process_document(content: bytes, file_type: str, url: str) -> Tuple[str, str, dict]:
        """
        Process document and extract text based on file type
        
        Args:
            content: Document content as bytes
            file_type: Type of file (html, pdf, docx, txt, etc.)
            url: URL of the document
            
        Returns:
            Tuple of (title, extracted_text, metadata)
        """
        file_type = file_type.lower()
        
        # HTML and server-side web pages (they typically render as HTML)
        if file_type in ['html', 'htm', 'aspx', 'asp', 'jsp', 'jspx', 'php', 'ashx', 'asmx', 'cfm', 'xhtml']:
            return DocumentProcessor.extract_text_from_html(content, url)
        elif file_type == 'pdf':
            return DocumentProcessor.extract_text_from_pdf(content, url)
        elif file_type in ['docx', 'doc']:
            return DocumentProcessor.extract_text_from_docx(content, url)
        elif file_type in ['txt', 'text', 'md', 'markdown']:
            return DocumentProcessor.extract_text_from_text(content, url)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return url, "", {}


# Singleton instance
_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get singleton instance of DocumentProcessor"""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
