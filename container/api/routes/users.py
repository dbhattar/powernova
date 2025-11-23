"""
User profile and document management routes
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, BackgroundTasks
from sqlalchemy.orm import Session
from sqlalchemy import desc
from pydantic import BaseModel, Field
from typing import List, Optional
import logging
from datetime import datetime

from database import get_db
from services.auth import get_current_user
from services.azure_storage import AzureStorageService
from services.embedding_processor import process_document_embedding
from models import User, Document, DocumentType, DocumentStatus, DocumentScope, ConversationDocument

router = APIRouter()
logger = logging.getLogger(__name__)


# ============================================================================
# Pydantic Models
# ============================================================================

class UserProfileResponse(BaseModel):
    """User profile response"""
    id: int
    email: str
    username: str
    is_active: bool
    is_verified: bool
    created_at: str
    total_conversations: int
    total_documents: int
    total_messages: int
    
    class Config:
        from_attributes = True


class UserProfileUpdate(BaseModel):
    """User profile update request"""
    username: Optional[str] = Field(None, min_length=1, max_length=100)


class PasswordChangeRequest(BaseModel):
    """Password change request"""
    current_password: str = Field(..., min_length=6)
    new_password: str = Field(..., min_length=6)


class DocumentResponse(BaseModel):
    """Document response model"""
    id: int
    title: str
    url: str
    document_type: str
    document_scope: str
    file_size: Optional[int] = None
    blob_url: Optional[str] = None
    status: str
    chunk_count: Optional[int] = None
    embedding_generated: bool
    created_at: str
    conversation_id: Optional[int] = None  # If linked to a conversation
    conversation_title: Optional[str] = None  # Title of the conversation


# ============================================================================
# User Profile Routes
# ============================================================================

@router.get("/profile", response_model=UserProfileResponse)
async def get_user_profile(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get current user's profile information
    
    Returns user details along with statistics:
    - Total conversations
    - Total documents uploaded
    - Total messages sent
    """
    # Count user's conversations
    from models import Conversation
    total_conversations = db.query(Conversation).filter(
        Conversation.user_id == current_user.id
    ).count()
    
    # Count user's documents (both conversation and user-scoped)
    total_documents = db.query(Document).filter(
        Document.uploaded_by == current_user.id
    ).count()
    
    # Count user's messages
    from models import Message
    total_messages = db.query(Message).join(
        Conversation
    ).filter(
        Conversation.user_id == current_user.id,
        Message.role == "user"
    ).count()
    
    return UserProfileResponse(
        id=current_user.id,
        email=current_user.email,
        username=current_user.username,
        is_active=current_user.is_active,
        is_verified=current_user.is_verified,
        created_at=current_user.created_at.isoformat(),
        total_conversations=total_conversations,
        total_documents=total_documents,
        total_messages=total_messages
    )


@router.put("/profile", response_model=UserProfileResponse)
async def update_user_profile(
    profile_update: UserProfileUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Update current user's profile information
    
    Currently supports:
    - username: Update display name
    """
    if profile_update.username is not None:
        current_user.username = profile_update.username
    
    db.commit()
    db.refresh(current_user)
    
    # Recalculate statistics
    from models import Conversation, Message
    total_conversations = db.query(Conversation).filter(
        Conversation.user_id == current_user.id
    ).count()
    
    total_documents = db.query(Document).filter(
        Document.uploaded_by == current_user.id
    ).count()
    
    total_messages = db.query(Message).join(
        Conversation
    ).filter(
        Conversation.user_id == current_user.id,
        Message.role == "user"
    ).count()
    
    return UserProfileResponse(
        id=current_user.id,
        email=current_user.email,
        username=current_user.username,
        is_active=current_user.is_active,
        is_verified=current_user.is_verified,
        created_at=current_user.created_at.isoformat(),
        total_conversations=total_conversations,
        total_documents=total_documents,
        total_messages=total_messages
    )


@router.post("/profile/change-password")
async def change_password(
    request: PasswordChangeRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Change current user's password
    
    Requires:
    - current_password: For verification
    - new_password: New password (min 6 characters)
    """
    from services.auth import verify_password, get_password_hash
    
    # Verify current password
    if not verify_password(request.current_password, current_user.hashed_password):
        raise HTTPException(
            status_code=400,
            detail="Current password is incorrect"
        )
    
    # Update password
    current_user.hashed_password = get_password_hash(request.new_password)
    current_user.must_change_password = False
    
    db.commit()
    
    return {
        "message": "Password changed successfully",
        "must_change_password": False
    }


# ============================================================================
# User Document Management Routes
# ============================================================================

@router.get("/documents", response_model=List[DocumentResponse])
async def get_user_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
    scope: Optional[str] = None,  # Filter by scope: 'user', 'conversation', or None for all
    limit: int = 100,
    offset: int = 0
):
    """
    Get all documents uploaded by the current user
    
    This includes:
    - Documents uploaded to user's personal library (scope='user')
    - Documents uploaded to conversations (scope='conversation')
    
    Query params:
    - scope: Filter by document scope ('user', 'conversation', or None for all)
    - limit: Maximum number of documents to return (default 100)
    - offset: Number of documents to skip (default 0)
    """
    query = db.query(Document).filter(
        Document.uploaded_by == current_user.id
    )
    
    # Apply scope filter if provided
    if scope:
        if scope not in ['user', 'conversation']:
            raise HTTPException(
                status_code=400,
                detail="Invalid scope. Must be 'user' or 'conversation'"
            )
        query = query.filter(Document.document_scope == scope)
    
    # Order by most recent first
    query = query.order_by(desc(Document.created_at))
    
    # Apply pagination
    documents = query.limit(limit).offset(offset).all()
    
    # Build response with conversation info if applicable
    result = []
    for doc in documents:
        # Find conversation link if it exists
        conversation_id = None
        conversation_title = None
        
        if doc.document_scope == DocumentScope.CONVERSATION:
            conv_doc = db.query(ConversationDocument).filter(
                ConversationDocument.document_id == doc.id
            ).first()
            
            if conv_doc and conv_doc.conversation:
                conversation_id = conv_doc.conversation.id
                conversation_title = conv_doc.conversation.title
        
        result.append(DocumentResponse(
            id=doc.id,
            title=doc.title,
            url=doc.url,
            document_type=doc.document_type.value,
            document_scope=doc.document_scope.value,
            file_size=doc.file_size,
            blob_url=doc.blob_url,
            status=doc.status.value,
            chunk_count=doc.chunk_count,
            embedding_generated=doc.embedding_generated,
            created_at=doc.created_at.isoformat(),
            conversation_id=conversation_id,
            conversation_title=conversation_title
        ))
    
    return result


@router.post("/documents", status_code=201)
async def upload_document_to_library(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Upload a document to user's personal library
    
    Documents uploaded here have scope='user' and are available
    across ALL of the user's conversations.
    
    Supported file types: PDF, DOCX, TXT, MD
    Max file size: 10MB
    """
    # Validate file size (10MB max)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    file_content = await file.read()
    file_size = len(file_content)
    
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is 10MB, got {file_size / (1024*1024):.2f}MB"
        )
    
    # Validate file type
    allowed_extensions = {'.pdf', '.docx', '.txt', '.md'}
    file_ext = '.' + file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    
    if file_ext not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
        )
    
    # Map file extension to DocumentType
    doc_type_map = {
        '.pdf': DocumentType.PDF,
        '.docx': DocumentType.DOCX,
        '.txt': DocumentType.TEXT,
        '.md': DocumentType.MARKDOWN
    }
    doc_type = doc_type_map.get(file_ext, DocumentType.OTHER)
    
    # Upload to Azure Blob Storage
    logger.info(f"Uploading document to user library: {file.filename} for user {current_user.id}")
    storage_service = AzureStorageService()
    
    # Generate unique URL for user library
    import hashlib
    file_hash = hashlib.md5(file_content).hexdigest()
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    unique_url = f"user_library/{current_user.id}/{file_hash}_{timestamp}"
    
    try:
        file_path, blob_url, uploaded_size = storage_service.upload_document(
            content=file_content,
            url=unique_url,
            file_extension=file_ext.lstrip('.'),
            job_id=0,  # User upload
            content_type=None
        )
        logger.info(f"✅ Successfully uploaded to Azure: {blob_url}")
    except Exception as e:
        logger.error(f"❌ Failed to upload to Azure Storage: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload document: {str(e)}"
        )
    
    # Create document record with USER scope
    document = Document(
        url=blob_url,
        title=file.filename,
        document_type=doc_type,
        document_scope=DocumentScope.USER,  # USER scope - available across all conversations
        uploaded_by=current_user.id,
        file_path=file_path,
        blob_url=blob_url,
        file_size=file_size,
        status=DocumentStatus.PENDING,
        crawl_job_id=None
    )
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    logger.info(f"Created document record {document.id} with USER scope")
    
    # Extract text content based on file type
    document.status = DocumentStatus.PROCESSING
    db.commit()
    
    try:
        if doc_type == DocumentType.PDF:
            import pdfplumber
            from io import BytesIO
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                text_content = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                document.content = text_content
                logger.info(f"Extracted {len(text_content)} chars from PDF")
        
        elif doc_type == DocumentType.TEXT or doc_type == DocumentType.MARKDOWN:
            document.content = file_content.decode('utf-8', errors='ignore')
            logger.info(f"Extracted {len(document.content)} chars from text file")
        
        elif doc_type == DocumentType.DOCX:
            from docx import Document as DocxDocument
            from io import BytesIO
            docx_doc = DocxDocument(BytesIO(file_content))
            text_content = '\n'.join(paragraph.text for paragraph in docx_doc.paragraphs)
            document.content = text_content
            logger.info(f"Extracted {len(text_content)} chars from DOCX")
        
        else:
            document.content = ""
            logger.warning(f"Unsupported document type: {doc_type.value}")
        
        # Check if we have content to process
        if not document.content or len(document.content.strip()) < 50:
            logger.warning(f"Document {document.id} has insufficient content ({len(document.content or '')} chars)")
            document.status = DocumentStatus.COMPLETED
            document.error_message = "Document contains insufficient text content"
            db.commit()
        else:
            # Mark as completed (embedding generation will happen in background)
            document.status = DocumentStatus.COMPLETED
            db.commit()
            db.refresh(document)
            
            # Generate embeddings in background
            background_tasks.add_task(process_document_embedding, document.id, db)
            logger.info(f"Queued document {document.id} for embedding generation")
    
    except Exception as e:
        logger.error(f"Error processing document {document.id}: {e}", exc_info=True)
        document.status = DocumentStatus.FAILED
        document.error_message = str(e)
        db.commit()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process document: {str(e)}"
        )
    
    return {
        "id": document.id,
        "title": document.title,
        "document_type": document.document_type.value,
        "document_scope": document.document_scope.value,
        "status": document.status.value,
        "file_size": document.file_size,
        "blob_url": document.blob_url,
        "message": "Document uploaded successfully to your library"
    }


@router.delete("/documents/{document_id}", status_code=204)
async def delete_user_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Delete a document from user's library
    
    Only the user who uploaded the document can delete it.
    This will also remove the document from any conversations it's linked to.
    """
    # Get document and verify ownership
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.uploaded_by == current_user.id
    ).first()
    
    if not document:
        raise HTTPException(
            status_code=404,
            detail="Document not found or you don't have permission to delete it"
        )
    
    # Delete from Azure Storage
    try:
        storage_service = AzureStorageService()
        # Extract blob name from blob_url or file_path
        if document.file_path:
            storage_service.delete_document(document.file_path)
            logger.info(f"Deleted document from Azure Storage: {document.file_path}")
    except Exception as e:
        logger.warning(f"Failed to delete document from Azure Storage: {e}")
        # Continue with database deletion even if storage deletion fails
    
    # Delete document (CASCADE will handle related records)
    db.delete(document)
    db.commit()
    
    logger.info(f"User {current_user.id} deleted document {document_id}")
    
    return None


# ============================================================================
# Helper Routes
# ============================================================================

@router.get("/documents/stats")
async def get_document_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Get statistics about user's documents
    
    Returns:
    - total_documents: Total number of documents uploaded
    - by_scope: Breakdown by document scope
    - by_status: Breakdown by processing status
    - by_type: Breakdown by file type
    """
    documents = db.query(Document).filter(
        Document.uploaded_by == current_user.id
    ).all()
    
    # Calculate statistics
    by_scope = {}
    by_status = {}
    by_type = {}
    total_size = 0
    
    for doc in documents:
        # By scope
        scope = doc.document_scope.value
        by_scope[scope] = by_scope.get(scope, 0) + 1
        
        # By status
        status = doc.status.value
        by_status[status] = by_status.get(status, 0) + 1
        
        # By type
        doc_type = doc.document_type.value
        by_type[doc_type] = by_type.get(doc_type, 0) + 1
        
        # Total size
        if doc.file_size:
            total_size += doc.file_size
    
    return {
        "total_documents": len(documents),
        "total_size_bytes": total_size,
        "total_size_mb": round(total_size / (1024 * 1024), 2),
        "by_scope": by_scope,
        "by_status": by_status,
        "by_type": by_type
    }
